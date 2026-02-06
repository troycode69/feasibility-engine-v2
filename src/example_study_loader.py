"""
Example Study Loader
Loads reference feasibility studies to improve AI-generated report quality
Supports: PDF, Word (.docx), Text (.txt), Markdown (.md)
"""

import os
from pathlib import Path
from typing import List, Dict

# Optional PDF support
try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False
    print("Warning: PyPDF2 not installed - PDF example studies won't load. Run: pip install PyPDF2")


def load_example_studies(example_dir: str = None) -> List[Dict[str, str]]:
    """
    Load all example studies from the example_studies directory.

    Supports: PDF, Word (.docx), Text (.txt), Markdown (.md)

    Returns:
        List of dicts with 'filename', 'content', and 'type' keys
    """
    if example_dir is None:
        # Default to src/data/example_studies
        current_dir = Path(__file__).parent
        example_dir = current_dir / "data" / "example_studies"
    else:
        example_dir = Path(example_dir)

    if not example_dir.exists():
        return []

    examples = []

    # Load PDF files
    for pdf_file in example_dir.glob("*.pdf"):
        try:
            content = extract_text_from_pdf(pdf_file)
            if content:
                examples.append({
                    'filename': pdf_file.name,
                    'content': content,
                    'type': 'pdf'
                })
        except Exception as e:
            print(f"Warning: Could not load {pdf_file.name}: {e}")

    # Load Word documents (.docx)
    for docx_file in example_dir.glob("*.docx"):
        try:
            content = extract_text_from_docx(docx_file)
            if content:
                examples.append({
                    'filename': docx_file.name,
                    'content': content,
                    'type': 'docx'
                })
        except Exception as e:
            print(f"Warning: Could not load {docx_file.name}: {e}")

    # Load text files
    for text_file in example_dir.glob("*.txt"):
        try:
            with open(text_file, 'r', encoding='utf-8') as f:
                content = f.read()
                examples.append({
                    'filename': text_file.name,
                    'content': content,
                    'type': 'text'
                })
        except Exception as e:
            print(f"Warning: Could not load {text_file.name}: {e}")

    # Load markdown files
    for md_file in example_dir.glob("*.md"):
        if md_file.name == "README.md":
            continue
        try:
            with open(md_file, 'r', encoding='utf-8') as f:
                content = f.read()
                examples.append({
                    'filename': md_file.name,
                    'content': content,
                    'type': 'markdown'
                })
        except Exception as e:
            print(f"Warning: Could not load {md_file.name}: {e}")

    return examples


def extract_text_from_pdf(pdf_path: Path) -> str:
    """Extract text content from a PDF file."""
    if not HAS_PYPDF2:
        print(f"Skipping PDF {pdf_path.name} - PyPDF2 not installed")
        return ""

    try:
        text_content = []
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text_content.append(page.extract_text())

        return "\n\n".join(text_content)
    except Exception as e:
        print(f"Error extracting PDF text: {e}")
        return ""


def extract_text_from_docx(docx_path: Path) -> str:
    """Extract text content from a Word document (.docx)."""
    try:
        from docx import Document

        doc = Document(docx_path)
        text_content = []

        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))

        return "\n\n".join(text_content)
    except ImportError:
        print(f"Warning: python-docx not installed. Install with: pip install python-docx")
        return ""
    except Exception as e:
        print(f"Error extracting Word document text: {e}")
        return ""


def format_examples_for_prompt(examples: List[Dict[str, str]], max_examples: int = 3) -> str:
    """
    Format example studies as context for AI prompts.

    Args:
        examples: List of example study dicts
        max_examples: Maximum number of examples to include (to control token usage)

    Returns:
        Formatted string to include in AI prompt
    """
    if not examples:
        return ""

    # Limit to max_examples
    examples = examples[:max_examples]

    prompt_text = "\n\n# REFERENCE EXAMPLES\n\n"
    prompt_text += "Below are example feasibility study sections from previous reports. "
    prompt_text += "Use these as reference for writing style, depth of analysis, and formatting:\n\n"

    for i, example in enumerate(examples, 1):
        # Truncate very long examples (max 10000 chars each for richer context)
        content = example['content']
        if len(content) > 10000:
            content = content[:10000] + "\n\n[... content truncated for length ...]"

        prompt_text += f"---\n\n## Example {i}: {example['filename']}\n\n"
        prompt_text += content
        prompt_text += "\n\n"

    prompt_text += "---\n\n"
    prompt_text += "Now, using the style and depth demonstrated above, generate the requested section for the current project.\n\n"

    return prompt_text


def extract_example_insights(examples: List[Dict[str, str]]) -> Dict:
    """
    Extract structured patterns from example studies for AI learning.

    Analyzes example feasibility studies to identify:
    - Common terminology and phrases
    - Score-to-recommendation mappings
    - Metric thresholds used for classifications
    - Risk assessment frameworks

    Args:
        examples: List of example study dicts with 'content' key

    Returns:
        Dict with extracted insights for AI training
    """
    import re

    insights = {
        'writing_patterns': [],
        'metric_thresholds': {},
        'recommendation_triggers': [],
        'risk_framework': [],
        'terminology': set(),
        'score_mappings': []
    }

    # Industry-specific terminology to look for
    storage_terms = [
        'NRSF', 'NOI', 'DSCR', 'cap rate', 'IRR', 'NPV',
        'lease-up', 'absorption', 'stabilization', 'occupancy',
        'SF per capita', 'rate compression', 'undersupplied',
        'oversupplied', 'balanced market', 'trade area',
        'climate-controlled', 'non-climate', 'unit mix'
    ]

    for example in examples:
        content = example.get('content', '')

        # Extract terminology usage
        for term in storage_terms:
            if term.lower() in content.lower():
                insights['terminology'].add(term)

        # Extract score-to-recommendation patterns
        # Look for patterns like "score of 85/100 = Proceed"
        score_patterns = re.findall(
            r'score[s]?\s*(?:of\s*)?(\d+)(?:/100|\s*points?).*?(proceed|caution|no.?go|excellent|good|fair|weak|poor)',
            content, re.IGNORECASE
        )
        for score, rec in score_patterns:
            insights['score_mappings'].append({
                'score': int(score),
                'recommendation': rec.lower()
            })

        # Extract SF per capita thresholds
        sf_patterns = re.findall(
            r'(\d+\.?\d*)\s*(?:SF|square feet)\s*per\s*capita.*?(undersupplied|oversupplied|balanced|saturated|healthy)',
            content, re.IGNORECASE
        )
        for value, classification in sf_patterns:
            if 'sf_per_capita' not in insights['metric_thresholds']:
                insights['metric_thresholds']['sf_per_capita'] = []
            insights['metric_thresholds']['sf_per_capita'].append({
                'value': float(value),
                'classification': classification.lower()
            })

        # Extract occupancy thresholds
        occ_patterns = re.findall(
            r'(\d+)%?\s*(?:occupancy|occ).*?(strong|healthy|weak|concerning|stabilized)',
            content, re.IGNORECASE
        )
        for value, classification in occ_patterns:
            if 'occupancy' not in insights['metric_thresholds']:
                insights['metric_thresholds']['occupancy'] = []
            insights['metric_thresholds']['occupancy'].append({
                'value': int(value),
                'classification': classification.lower()
            })

        # Extract risk keywords
        risk_keywords = re.findall(
            r'(risk|concern|challenge|threat|weakness)[:;]\s*([^.]+\.)',
            content, re.IGNORECASE
        )
        for risk_type, description in risk_keywords:
            insights['risk_framework'].append({
                'type': risk_type.lower(),
                'description': description.strip()
            })

        # Extract recommendation triggers
        rec_patterns = re.findall(
            r'(?:recommend|suggests?|indicates?)\s+(?:to\s+)?(proceed|caution|not\s+proceed|pass|investigate)',
            content, re.IGNORECASE
        )
        insights['recommendation_triggers'].extend([r.lower() for r in rec_patterns])

    # Convert set to list for JSON serialization
    insights['terminology'] = list(insights['terminology'])

    # Deduplicate
    insights['recommendation_triggers'] = list(set(insights['recommendation_triggers']))

    return insights


def get_style_guide_from_examples(examples: List[Dict[str, str]]) -> str:
    """
    Generate a concise style guide from example studies.

    Returns a formatted string that can be prepended to AI prompts
    to ensure consistent writing style.

    Args:
        examples: List of example study dicts

    Returns:
        Style guide text
    """
    insights = extract_example_insights(examples)

    style_guide = "=== STYLE GUIDE FROM REFERENCE STUDIES ===\n\n"

    # Terminology
    if insights['terminology']:
        style_guide += "**Industry Terminology to Use:**\n"
        style_guide += ", ".join(sorted(insights['terminology']))
        style_guide += "\n\n"

    # Score mappings
    if insights['score_mappings']:
        style_guide += "**Score-to-Recommendation Framework:**\n"
        for mapping in insights['score_mappings'][:5]:  # Limit to 5
            style_guide += f"  - Score {mapping['score']}: {mapping['recommendation'].title()}\n"
        style_guide += "\n"

    # Metric thresholds
    if insights['metric_thresholds'].get('sf_per_capita'):
        style_guide += "**SF Per Capita Benchmarks:**\n"
        for threshold in insights['metric_thresholds']['sf_per_capita'][:3]:
            style_guide += f"  - {threshold['value']} SF/capita = {threshold['classification'].title()}\n"
        style_guide += "\n"

    # Risk patterns
    if insights['risk_framework']:
        style_guide += "**Risk Assessment Patterns:**\n"
        for risk in insights['risk_framework'][:3]:
            style_guide += f"  - {risk['type'].title()}: {risk['description'][:80]}...\n"
        style_guide += "\n"

    style_guide += "==========================================\n\n"

    return style_guide


# Test function
if __name__ == "__main__":
    examples = load_example_studies()
    print(f"Loaded {len(examples)} example studies:")
    for ex in examples:
        print(f"  - {ex['filename']} ({ex['type']}, {len(ex['content'])} chars)")

    if examples:
        prompt_addition = format_examples_for_prompt(examples, max_examples=2)
        print(f"\nPrompt addition length: {len(prompt_addition)} characters")

        # Test insight extraction
        print("\n=== Testing Insight Extraction ===")
        insights = extract_example_insights(examples)
        print(f"Terminology found: {len(insights['terminology'])} terms")
        print(f"Score mappings: {len(insights['score_mappings'])}")
        print(f"Metric thresholds: {list(insights['metric_thresholds'].keys())}")
        print(f"Risk patterns: {len(insights['risk_framework'])}")

        # Test style guide generation
        print("\n=== Style Guide ===")
        style_guide = get_style_guide_from_examples(examples)
        print(style_guide[:500] + "...")
